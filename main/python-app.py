import sys
import streamlit as st
import pandas as pd
import random
import io
# from openai import OpenAI # Not directly used for OpenRouter
import json
from io import BytesIO
from typing import List, Optional, Dict, Any, Tuple, AsyncGenerator
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from pdfminer.high_level import extract_text as pdf_extract_text
import tempfile
import time
# import openai # Not directly used for OpenRouter
import requests
import os
import numpy as np
from sklearn.metrics import cohen_kappa_score
from itertools import combinations
import asyncio
import aiohttp
from functools import lru_cache
import hashlib
import logging
from openai import OpenAI

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize session state variables
if "analysis_mode" not in st.session_state:
    st.session_state.analysis_mode = "Inductive"
if "transcript_text" not in st.session_state:
    st.session_state.transcript_text = None
if "research_question" not in st.session_state:
    st.session_state.research_question = ""
if "framework" not in st.session_state:
    st.session_state.framework = None
if "analysis_complete" not in st.session_state:
    st.session_state.analysis_complete = False
if "analysis_results" not in st.session_state:
    st.session_state.analysis_results = {}
if "coded_segments" not in st.session_state:
    st.session_state.coded_segments = []
if "themes" not in st.session_state:
    st.session_state.themes = []
if "analysis_start_time" not in st.session_state:
    st.session_state.analysis_start_time = None
if "coded_results" not in st.session_state:
    st.session_state.coded_results = []

# Constants
VALID_TRANSCRIPT_TYPES = [".txt", ".docx", ".pdf"]
VALID_FRAMEWORK_TYPES = [".txt", ".docx", ".pdf"]  # Updated framework file types
API_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = "deepseek/deepseek-chat:free"
MAX_TOKENS = 4000
TEMPERATURE = 0.7
CHUNK_SIZE = 1500
OVERLAP = 200
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB
CHUNK_READ_SIZE = 1024 * 1024  # 1 MB chunks for streaming

# Theme colors for highlighting
THEME_COLORS = {
    "Learning Experience": "#FFB6C1",  # Light pink
    "User Interface": "#98FB98",       # Light green
    "Technical Issues": "#87CEEB",     # Sky blue
    "Feedback": "#DDA0DD",             # Plum
    "General": "#F0E68C"               # Khaki
}

# Get API key from secrets
API_KEY = st.secrets["OPENROUTER_API_KEY"]

# Define the headers for the API request
headers = {
    "Authorization": f"Bearer {API_KEY}",
    "Content-Type": "application/json",
    "HTTP-Referer": "https://github.com/baha-ti/Qualinsight-AI",
    "X-Title": "Qualinsight AI"
}

@st.cache_data(ttl=3600)  # Cache for 1 hour
def extract_text_from_txt(file) -> str:
    """Extract text from TXT file with caching."""
    logger.info("Extracting text from TXT file (cache miss)")
    return file.getvalue().decode('utf-8')

@st.cache_data(ttl=3600)
def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file with caching."""
    logger.info("Extracting text from DOCX file (cache miss)")
    doc = Document(file)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

@st.cache_data(ttl=3600)
def extract_text_from_pdf(file) -> str:
    """Extract text from PDF file with caching."""
    logger.info("Extracting text from PDF file (cache miss)")
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
        tmp_file.write(file.getvalue())
        tmp_file.flush()
        text = pdf_extract_text(tmp_file.name)
    os.unlink(tmp_file.name)
    return text

@st.cache_data(ttl=3600)
def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = OVERLAP) -> List[str]:
    """Split text into overlapping chunks with context preservation and caching."""
    logger.info("Chunking text (cache miss)")
    words = text.split()
    chunks = []
    start = 0
    
    while start < len(words):
        # Get chunk of words
        end = min(start + chunk_size, len(words))
        chunk = ' '.join(words[start:end])
        
        # Add context from previous chunk if not first chunk
        if start > 0:
            context_start = max(0, start - overlap)
            context = ' '.join(words[context_start:start])
            chunk = f"{context}\n\n{chunk}"
        
        # Add context from next chunk if not last chunk
        if end < len(words):
            context_end = min(len(words), end + overlap)
            context = ' '.join(words[end:context_end])
            chunk = f"{chunk}\n\n{context}"
        
        chunks.append(chunk)
        start = end
    
    return chunks

async def stream_file(file, chunk_size: int = CHUNK_READ_SIZE) -> AsyncGenerator[bytes, None]:
    """Stream file in chunks."""
    while True:
        chunk = file.read(chunk_size)
        if not chunk:
            break
        yield chunk

async def process_large_file(file) -> str:
    """Process large files in chunks."""
    content = []
    async for chunk in stream_file(file):
        content.append(chunk)
    return b''.join(content).decode('utf-8')

async def get_ai_response_async(messages: List[Dict], stream: bool = True) -> AsyncGenerator[str, None]:
    """Get response from DeepSeek API asynchronously."""
    logger.info("Starting async API request")
    start_time = time.time()
    
    data = {
        "model": MODEL,
        "messages": messages,
        "max_tokens": MAX_TOKENS,
        "temperature": TEMPERATURE,
        "stream": stream
    }
    
    async with aiohttp.ClientSession() as session:
        async with session.post(API_URL, json=data, headers=headers) as response:
            if response.status == 200:
                if stream:
                    async for line in response.content:
                        if line:
                            try:
                                json_response = json.loads(line.decode('utf-8').replace('data: ', ''))
                                if 'choices' in json_response and len(json_response['choices']) > 0:
                                    content = json_response['choices'][0].get('delta', {}).get('content', '')
                                    if content:
                                        yield content
                            except json.JSONDecodeError:
                                continue
                else:
                    result = await response.json()
                    yield result['choices'][0]['message']['content']
                
                logger.info(f"API request completed in {time.time() - start_time:.2f} seconds")
            else:
                error_text = await response.text()
                logger.error(f"API Error: {response.status} - {error_text}")
                raise Exception(f"API Error: {response.status} - {error_text}")

def get_ai_response(messages: List[Dict], stream: bool = True):
    """Synchronous wrapper for async AI response."""
    if stream:
        return asyncio.run(get_ai_response_async(messages, stream))
    else:
        return asyncio.run(get_ai_response_async(messages, stream))

def process_transcript(uploaded_file) -> Optional[List[str]]:
    """Process uploaded transcript files with size limits and streaming."""
    start_time = time.time()
    logger.info(f"Processing file: {uploaded_file.name} ({uploaded_file.size/1024/1024:.2f}MB)")
    
    if uploaded_file.size > MAX_FILE_SIZE:
        st.warning(f"File size exceeds {MAX_FILE_SIZE/1024/1024}MB limit. Processing in chunks...")
        try:
            # Process large file asynchronously
            text_content = asyncio.run(process_large_file(uploaded_file))
            logger.info(f"Large file processed in {time.time() - start_time:.2f} seconds")
        except Exception as e:
            logger.error(f"Error processing large file: {str(e)}")
            st.error(f"Error processing large file: {str(e)}")
            return None
    else:
        file_type = uploaded_file.name.split('.')[-1].lower()
        try:
            if file_type == "txt":
                text_content = extract_text_from_txt(uploaded_file)
            elif file_type == "docx":
                text_content = extract_text_from_docx(uploaded_file)
            elif file_type == "pdf":
                text_content = extract_text_from_pdf(uploaded_file)
            else:
                st.error(f"Unsupported file type: {file_type}")
                return None
            logger.info(f"File processed in {time.time() - start_time:.2f} seconds")
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            st.error(f"Error extracting text: {str(e)}")
            return None

    if text_content:
        # Chunk the text with caching
        chunks = chunk_text(text_content)
        logger.info(f"Text chunked into {len(chunks)} chunks")
        return chunks
    return None

def highlight_text(text, theme, color):
    """Wrap text in HTML mark tag with specified color"""
    return f'<mark style="background-color: {color}">{text}</mark>'

# Initialize the Streamlit app
st.set_page_config(
    page_title="Qualinsight AI - Qualitative Analysis",
    page_icon="🔍",
    layout="wide"
)

# Knowledge Base File Operations
def load_knowledge_base() -> Dict[str, Dict[str, Any]]:
    """Load knowledge base from JSON file"""
    kb_path = os.path.join(os.path.dirname(__file__), 'knowledge_base.json')
    if os.path.exists(kb_path):
        with open(kb_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_knowledge_base(kb: Dict[str, Dict[str, Any]]):
    """Save knowledge base to JSON file"""
    kb_path = os.path.join(os.path.dirname(__file__), 'knowledge_base.json')
    with open(kb_path, 'w', encoding='utf-8') as f:
        json.dump(kb, f, ensure_ascii=False, indent=2)

# Load knowledge base at startup
st.session_state.knowledge_base = load_knowledge_base()

# --- Title ---
st.title("QualInsight AI - Qualitative Research Assistant")

# Create tabs for different sections
tab1, tab2, tab3, tab4 = st.tabs(["Input", "Analysis", "Results", "Export"])

with tab1:
    st.header("Input")
    
    # Input method selection
    input_method = st.radio(
        "Choose input method:",
        ["Upload File", "Paste Text"],
        key="input_method"
    )
    
    # File upload or text input (moved before research question)
    if input_method == "Upload File":
        uploaded_file = st.file_uploader(
            "Upload transcript (.txt, .docx, or .pdf)",
            type=VALID_TRANSCRIPT_TYPES,
            key="transcript_uploader"
        )
        if uploaded_file is not None:
            if uploaded_file.size > MAX_FILE_SIZE:
                st.warning(f"File size ({uploaded_file.size/1024/1024:.1f}MB) exceeds {MAX_FILE_SIZE/1024/1024}MB limit. Processing may take longer.")
            st.session_state.transcript_text = process_transcript(uploaded_file)
    else:
        st.session_state.transcript_text = st.text_area(
            "Paste your transcript here:",
            value=st.session_state.transcript_text if st.session_state.transcript_text else "",
            key="transcript_text_input"
        )
    
    # Research question input
    st.session_state.research_question = st.text_area(
        "Enter your research question:",
        value=st.session_state.research_question,
        key="research_question_input"
    )
    
    # Analysis mode selection
    st.session_state.analysis_mode = st.radio(
        "Choose analysis mode:",
        ["Inductive", "Deductive"],
        key="analysis_mode_radio"
    )
    
    # Framework upload for deductive analysis
    if st.session_state.analysis_mode == "Deductive":
        framework_file = st.file_uploader(
            "Upload theoretical framework (PDF, DOCX, or TXT)",
            type=VALID_FRAMEWORK_TYPES,
            key="framework_uploader",
            help="Upload your theoretical framework document that guides your research analysis"
        )
        if framework_file is not None:
            try:
                if framework_file.name.endswith('.pdf'):
                    st.session_state.framework = extract_text_from_pdf(framework_file)
                elif framework_file.name.endswith('.docx'):
                    st.session_state.framework = extract_text_from_docx(framework_file)
                else:  # .txt
                    st.session_state.framework = extract_text_from_txt(framework_file)
                st.success("Framework loaded successfully!")
            except Exception as e:
                st.error(f"Error loading framework: {str(e)}")

with tab2:
    st.header("Analysis")
    
    # Analysis parameters
    with st.expander("Analysis Parameters"):
        temperature = st.slider(
            "Temperature (creativity)",
            min_value=0.0,
            max_value=1.0,
            value=TEMPERATURE,
            step=0.1,
            key="temperature_slider"
        )
        max_tokens = st.slider(
            "Max Tokens",
            min_value=1000,
            max_value=8000,
            value=MAX_TOKENS,
            step=1000,
            key="max_tokens_slider"
        )
    
    # Start analysis button
    if st.button("Start Analysis", key="start_analysis"):
        st.info("Start Analysis button clicked!")
        if not st.session_state.transcript_text:
            st.error("Please provide a transcript first.")
            st.info(f"Transcript text: {st.session_state.transcript_text}")
        elif not st.session_state.research_question:
            st.error("Please enter a research question.")
            st.info(f"Research question: {st.session_state.research_question}")
        elif st.session_state.analysis_mode == "Deductive" and not st.session_state.framework:
            st.error("Please upload a theoretical framework for deductive analysis.")
            st.info(f"Analysis mode: {st.session_state.analysis_mode}, Framework: {st.session_state.framework}")
        else:
            st.info("All conditions met, starting analysis...")
            st.session_state.analysis_start_time = time.time()
            st.session_state.analysis_complete = False
            st.session_state.coded_segments = []
            st.session_state.themes = []
            
            # Create progress bar
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                # Process the transcript
                if isinstance(st.session_state.transcript_text, list):
                    st.session_state.total_chunks = len(st.session_state.transcript_text)
                    st.info(f"Processing {st.session_state.total_chunks} chunks...")
                    for i, chunk in enumerate(st.session_state.transcript_text):
                        st.session_state.current_chunk = i + 1
                        status_text.text(f"Processing chunk {i+1} of {st.session_state.total_chunks}")
                        progress_bar.progress((i + 1) / st.session_state.total_chunks)
                        
                        # Analyze chunk with AI
                        with st.spinner("🤖 AI is analyzing chunk {i+1} of {st.session_state.total_chunks}..."):
                            analysis_result = analyze_transcript(
                                chunk,
                                st.session_state.research_question,
                                st.session_state.analysis_mode,
                                st.session_state.framework
                            )
                        
                        if analysis_result:
                            st.session_state.coded_segments.extend(analysis_result.get('coded_segments', []))
                            st.session_state.themes.extend(analysis_result.get('themes', []))
                            st.info(f"Chunk {i+1} analysis result processed.")
                
                # After processing all chunks, set coded_results
                st.session_state.coded_results = st.session_state.coded_segments
                st.info(f"Coded results set: {len(st.session_state.coded_results)} segments.")

                st.session_state.analysis_complete = True
                status_text.text("Analysis complete!")
                progress_bar.progress(1.0)
                
                # Show results immediately
                st.success("Analysis completed!")
                st.rerun()  # Refresh to show results
                
            except Exception as e:
                st.error(f"Error during analysis: {str(e)}")
                status_text.text("Analysis failed!")
                progress_bar.progress(0)

with tab3:
    st.header("Step 3: Results")
    
    if 'coded_results' in st.session_state and st.session_state.coded_results:
        # Create DataFrame for initial coding
        df_initial = pd.DataFrame(st.session_state.coded_results)
        
        # Display results in expandable sections
        with st.expander("Coded Segments", expanded=True):
            st.data_editor(
                df_initial[['text', 'code', 'notes']],
                key="coded_segments_editor",
                num_rows="dynamic",
                use_container_width=True
            )
            
            # Update the original dataframe with edited values
            # The data_editor directly modifies the dataframe, so we just need to ensure session state is updated
            st.session_state.coded_results = df_initial.to_dict('records')
                
            st.write("---")
        
        # Theme Development
        with st.expander("Theme Development", expanded=False):
            if st.button("Proceed to Theme Development", key="theme_development"):
                try:
                    with st.spinner("🤖 AI is developing themes and subthemes... This may take a few minutes."):
                        # Second stage prompt - focus on theme development
                        system_prompt = """You are an AI assistant developing themes from coded transcript segments.
                        For this second stage, focus on:
                        1. Grouping related codes into themes
                        2. Identifying subthemes within each theme
                        3. Providing evidence from the transcript

                        IMPORTANT: Your response MUST be a plain text string, with each theme entry separated by `###THEME_ENTRY###`. Each entry must contain the following fields, delimited by `|||`:
                        - THEME: <main theme>
                        - SUBTHEME: <subtheme within the main theme>
                        - CODES: <list of related codes, comma-separated>
                        - EVIDENCE: <list of relevant quotes, comma-separated>
                        - EXPLANATION: <brief explanation of this theme>

                        Example response format:
                        THEME: Learning Experience ||| SUBTHEME: Instructional Clarity ||| CODES: Clear Instructions, Understanding ||| EVIDENCE: The instructions were very clear, I understood what to do, it was "spot on" ||| EXPLANATION: Participants found the instructions clear and easy to follow###THEME_ENTRY###
                        """
                        
                        # Convert initial coding to string for analysis
                        coding_summary_str = df_initial.to_string(index=False)
                        
                        messages = [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": f"Research Questions: {st.session_state.research_question}\n\nCoded Transcript:\n{coding_summary_str}"}
                        ]
                        
                        response = get_ai_response(messages)
                        if response:
                            try:
                                theme_entries = response.split('###THEME_ENTRY###')
                                theme_results = []
                                
                                for entry in theme_entries:
                                    if entry.strip():
                                        parts = entry.split('|||')
                                        if len(parts) == 5:
                                            theme_result = {
                                                'theme': parts[0].replace('THEME:', '').strip(),
                                                'subtheme': parts[1].replace('SUBTHEME:', '').strip(),
                                                'codes': [code.strip() for code in parts[2].replace('CODES:', '').split(',')],
                                                'evidence': [quote.strip().strip('"\'') for quote in parts[3].replace('EVIDENCE:', '').split(',')],
                                                'explanation': parts[4].replace('EXPLANATION:', '').strip()
                                            }
                                            theme_results.append(theme_result)
                                
                                df_themes = pd.DataFrame(theme_results)
                                st.session_state.theme_results = df_themes
                                st.dataframe(df_themes)
                                
                                # Create highlighted transcript
                                with st.expander("Highlighted Transcript", expanded=False):
                                    highlighted_text = st.session_state.transcript_text[0]  # Assuming single transcript
                                    
                                    # Apply highlights
                                    for _, row in df_themes.iterrows():
                                        for quote in row['evidence']:
                                            if quote in highlighted_text:
                                                theme = row['theme']
                                                color = THEME_COLORS.get(theme, THEME_COLORS["General"])
                                                highlighted_text = highlighted_text.replace(
                                                    quote,
                                                    highlight_text(quote, theme, color)
                                                )
                                    
                                    st.markdown(highlighted_text, unsafe_allow_html=True)
                                
                            except Exception as e:
                                st.error(f"Failed to parse theme development response. Error: {str(e)}")
                                st.text("Raw response:")
                                st.text(response)
                except Exception as e:
                    st.error(f"Error in theme development: {str(e)}")
        
        # Inter-coder Reliability
        with st.expander("Inter-coder Reliability", expanded=False):
            st.write("Upload coding results from multiple coders to calculate reliability metrics")
            
            uploaded_files = st.file_uploader(
                "Upload coding results (CSV files)",
                type=["csv"],
                accept_multiple_files=True
            )
            
            if uploaded_files and len(uploaded_files) >= 2:
                try:
                    coders_data = []
                    for file in uploaded_files:
                        df = pd.read_csv(file)
                        coders_data.append(df)
                    
                    reliability_scores = calculate_intercoder_reliability(coders_data)
                    
                    st.write("### Reliability Scores")
                    for pair, score in reliability_scores.items():
                        st.write(f"{pair}: {score:.3f}")
                    
                except Exception as e:
                    st.error(f"Error calculating reliability: {str(e)}")

with tab4:
    st.header("Step 4: Export")
    
    if 'coded_results' in st.session_state and 'theme_results' in st.session_state:
        # Export Options
        st.subheader("Export Options")
        
        # Export initial coding
        csv_initial = pd.DataFrame(st.session_state.coded_results).to_csv(index=False)
        st.download_button(
            "Download Initial Coding (CSV)",
            csv_initial,
            "initial_coding.csv",
            "text/csv",
            key="download_initial_csv"
        )
        
        # Export themes
        csv_themes = st.session_state.theme_results.to_csv(index=False)
        st.download_button(
            "Download Themes (CSV)",
            csv_themes,
            "themes.csv",
            "text/csv",
            key="download_themes_csv"
        )
        
        # Export full report
        if st.button("Generate Full Report"):
            try:
                # Create Word document
                doc = Document()
                doc.add_heading("Qualitative Research Analysis Report", 0)
                
                # Add research questions
                doc.add_heading("Research Questions", level=1)
                doc.add_paragraph(st.session_state.research_question)
                
                # Add initial coding
                doc.add_heading("Initial Coding", level=1)
                for respondent in pd.DataFrame(st.session_state.coded_results)['respondent'].unique():
                    doc.add_heading(f"Respondent: {respondent}", level=2)
                    respondent_df = pd.DataFrame(st.session_state.coded_results)
                    respondent_df = respondent_df[respondent_df['respondent'] == respondent]
                    for _, row in respondent_df.iterrows():
                        doc.add_paragraph(f"Code: {row['code']}")
                        doc.add_paragraph(f"Text: {row['text']}")
                        doc.add_paragraph(f"Notes: {row['notes']}")
                        doc.add_paragraph("---")
                
                # Add themes
                doc.add_heading("Theme Development", level=1)
                for _, row in st.session_state.theme_results.iterrows():
                    doc.add_heading(f"Theme: {row['theme']}", level=2)
                    doc.add_paragraph(f"Subtheme: {row['subtheme']}")
                    doc.add_paragraph(f"Codes: {', '.join(row['codes'])}")
                    doc.add_paragraph(f"Evidence: {', '.join(row['evidence'])}")
                    doc.add_paragraph(f"Explanation: {row['explanation']}")
                    doc.add_paragraph("---")
                
                # Add highlighted transcript
                doc.add_heading("Highlighted Transcript", level=1)
                doc.add_paragraph(st.session_state.transcript_text[0])
                
                doc_stream = BytesIO()
                doc.save(doc_stream)
                doc_stream.seek(0)
                st.download_button(
                    "Download Full Report (Word)",
                    doc_stream,
                    "analysis_report.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_word"
                )
                
            except Exception as e:
                st.error(f"Error generating report: {str(e)}")

def merge_coded_chunks(chunks: List[Dict]) -> List[Dict]:
    """Merge coded chunks while handling overlapping segments."""
    merged = []
    seen_texts = set()
    
    for chunk in chunks:
        for entry in chunk:
            # Skip if we've seen this exact text before
            if entry['text'] in seen_texts:
                continue
            
            # Check for overlapping segments
            is_overlap = False
            for existing in merged:
                if entry['text'] in existing['text'] or existing['text'] in entry['text']:
                    # Keep the longer segment
                    if len(entry['text']) > len(existing['text']):
                        merged.remove(existing)
                        merged.append(entry)
                    is_overlap = True
                    break
            
            if not is_overlap:
                merged.append(entry)
                seen_texts.add(entry['text'])
    
    return merged

def calculate_intercoder_reliability(coders_data: List[pd.DataFrame]) -> Dict[str, float]:
    """Calculate Cohen's Kappa and Fleiss' Kappa for multiple coders."""
    if len(coders_data) < 2:
        return {"error": "Need at least 2 coders for reliability analysis"}
    
    # Get unique segments across all coders
    all_segments = set()
    for df in coders_data:
        all_segments.update(df['text'].tolist())
    
    # Create coding matrices
    segment_codes = {segment: [] for segment in all_segments}
    
    for df in coders_data:
        for _, row in df.iterrows():
            segment_codes[row['text']].append(row['code'])
    
    # Calculate pairwise Cohen's Kappa
    kappa_scores = {}
    for i, j in combinations(range(len(coders_data)), 2):
        coder1_codes = [codes[i] if i < len(codes) else None for codes in segment_codes.values()]
        coder2_codes = [codes[j] if j < len(codes) else None for codes in segment_codes.values()]
        
        # Filter out None values
        valid_pairs = [(c1, c2) for c1, c2 in zip(coder1_codes, coder2_codes) if c1 is not None and c2 is not None]
        if valid_pairs:
            c1_codes, c2_codes = zip(*valid_pairs)
            kappa = cohen_kappa_score(c1_codes, c2_codes)
            kappa_scores[f"Coders {i+1}-{j+1}"] = kappa
    
    return kappa_scores

def load_framework(file) -> Dict:
    """Load and validate a coding framework from JSON file."""
    try:
        framework = json.load(file)
        required_keys = ['name', 'description', 'codes', 'categories']
        if not all(key in framework for key in required_keys):
            raise ValueError("Framework missing required keys")
        return framework
    except json.JSONDecodeError:
        raise ValueError("Invalid JSON format")
    except Exception as e:
        raise ValueError(f"Error loading framework: {str(e)}")

def ai_generate_codes(text: str, mode: str, rq: str, framework: Optional[str] = None) -> List[dict]:
    """Generates initial codes or themes using the AI."""
    # This function is not used directly anymore, but is kept for reference or future use.
    # The prompt building and API call logic is now directly in the main app flow for more control.
    pass

def analyze_transcript(transcript_text: str, research_question: str, analysis_mode: str, framework: Optional[Dict] = None) -> Dict:
    """Analyze transcript using OpenAI API with streaming response"""
    try:
        # Create progress container
        progress_container = st.empty()
        progress_container.info("🔄 Initializing AI analysis...")
        
        # Prepare the prompt
        prompt = f"""Analyze the following transcript in the context of this research question: "{research_question}"

Transcript:
{transcript_text}

Please provide a detailed analysis including:
1. Key themes and patterns
2. Relevant quotes and examples
3. Insights and implications
4. Recommendations for further research

Format the response in clear sections with markdown formatting."""

        # Initialize OpenAI client
        client = OpenAI(api_key=st.secrets["openai"]["api_key"])
        
        # Update progress
        progress_container.info("🤖 AI is processing the content...")
        
        # Get streaming response
        stream = client.chat.completions.create(
            model="gpt-4-turbo-preview",
            messages=[{"role": "user", "content": prompt}],
            stream=True
        )
        
        # Process streaming response
        full_response = ""
        response_container = st.empty()
        
        for chunk in stream:
            if chunk.choices[0].delta.content:
                full_response += chunk.choices[0].delta.content
                response_container.markdown(full_response)
        
        # Update progress
        progress_container.info("✨ Analysis complete! Processing results...")
        
        # Process and structure the response
        # Split the response into sections
        sections = full_response.split('\n\n')
        themes = []
        quotes = []
        insights = []
        recommendations = []
        
        current_section = None
        for section in sections:
            if section.startswith('1.') or 'themes' in section.lower():
                current_section = 'themes'
            elif section.startswith('2.') or 'quotes' in section.lower():
                current_section = 'quotes'
            elif section.startswith('3.') or 'insights' in section.lower():
                current_section = 'insights'
            elif section.startswith('4.') or 'recommendations' in section.lower():
                current_section = 'recommendations'
            
            if current_section == 'themes':
                themes.append(section)
            elif current_section == 'quotes':
                quotes.append(section)
            elif current_section == 'insights':
                insights.append(section)
            elif current_section == 'recommendations':
                recommendations.append(section)
        
        analysis_result = {
            "themes": themes,
            "quotes": quotes,
            "insights": insights,
            "recommendations": recommendations,
            "coded_segments": [{"text": q, "code": "Auto-coded", "notes": ""} for q in quotes]
        }
        
        # Clear progress message
        progress_container.empty()
        
        return analysis_result
        
    except Exception as e:
        st.error(f"Error during analysis: {str(e)}")
        return {}

# Add debug information to the UI
if st.sidebar.checkbox("Show Debug Information"):
    st.sidebar.write("### Debug Information")
    st.sidebar.write(f"Cache TTL: 3600 seconds")
    st.sidebar.write(f"Max File Size: {MAX_FILE_SIZE/1024/1024}MB")
    st.sidebar.write(f"Chunk Size: {CHUNK_SIZE} words")
    st.sidebar.write(f"Chunk Overlap: {OVERLAP} words")
    
    if 'analysis_start_time' in st.session_state:
        st.sidebar.write(f"Last Analysis Duration: {time.time() - st.session_state.analysis_start_time:.2f} seconds")