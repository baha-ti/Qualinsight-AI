import sys
import streamlit as st
import pandas as pd
import re
import random
import io
import openai
import json
from io import BytesIO
from typing import List, Optional
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from pdfminer.high_level import extract_text as pdf_extract_text
import tempfile

# Constants
VALID_TRANSCRIPT_TYPES = ["txt", "docx", "pdf"]
VALID_FRAMEWORK_TYPES = ["json", "txt"]
OPENAI_MODEL = "gpt-4"
MAX_TOKENS = 1500
TEMPERATURE = 0.5

# --- Title ---
st.title("QualInsight AI - Qualitative Research Assistant")
st.markdown("Upload your transcript, frameworks, and let the AI generate codes, themes, highlights, and reports.")

# --- Sidebar for API Key ---
st.sidebar.header("Configuration")
openai_key = st.sidebar.text_input("OpenAI API Key", type="password")

st.sidebar.info(f"Python executable: {sys.executable}")

if openai_key and openai_key.startswith("sk-"):
    openai.api_key = openai_key
else:
    if openai_key:
        st.sidebar.warning("Invalid key format. Make sure it starts with 'sk-'.")
    else:
        st.sidebar.info("Enter your OpenAI API Key to enable AI integration.")

# --- Upload section ---
uploaded_file = st.file_uploader("Upload transcript (.txt, .docx, or .pdf)", type=VALID_TRANSCRIPT_TYPES)
framework_file = st.file_uploader("Upload framework for deductive coding (.json or .txt)", type=VALID_FRAMEWORK_TYPES)

# --- Research question input ---
research_questions = st.text_area("Enter your research question(s)", height=100)

# --- Analysis mode ---
analysis_mode = st.selectbox("Choose coding mode", ["Inductive", "Deductive", "Hybrid"])

# --- Extract text functions ---
def extract_text_from_txt(file) -> str:
    try:
        return file.read().decode("utf-8")
    except UnicodeDecodeError:
        file.seek(0)
        return file.read().decode("latin-1")


def extract_text_from_docx(file) -> str:
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])


def extract_text_from_pdf(file) -> str:
    with tempfile.NamedTemporaryFile(delete=True, suffix=".pdf") as temp_file:
        temp_file.write(file.read())
        temp_file.seek(0)
        return pdf_extract_text(temp_file.name)


def extract_text(file) -> str:
    if not file:
        return ""
    try:
        file.seek(0)
        ext = file.name.split('.')[-1].lower()
        if ext == 'txt':
            return extract_text_from_txt(file)
        elif ext == 'docx':
            return extract_text_from_docx(file)
        elif ext == 'pdf':
            return extract_text_from_pdf(file)
        else:
            st.error(f"Unsupported file type: {ext}")
            return ""
    except Exception as e:
        st.error(f"Failed to extract text: {e}")
        return ""

# --- AI-based code generator ---
def ai_generate_codes(text: str, mode: str, rq: str, framework: Optional[str] = None) -> List[dict]:
    if not openai.api_key:
        st.error("Please configure a valid OpenAI API Key in the sidebar.")
        return []

    system_msg = (
        f"You are a qualitative research assistant trained in thematic analysis. Apply mode: {mode}."
    )
    if framework:
        system_msg += f" Use this framework for deductive coding: {framework}"

    try:
        response = openai.ChatCompletion.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": (
                    f"Research Questions:\n{rq}\n\n"
                    f"Framework:\n{framework or 'None'}\n\n"
                    f"Transcript:\n{text}"
                )}
            ],
            temperature=TEMPERATURE,
            max_tokens=MAX_TOKENS
        )
        content = response.choices[0].message.content.strip()
        data = json.loads(content)
        if not isinstance(data, list):
            st.error("AI output is not a list. Check your prompt and input.")
            return []
        return data
    except json.JSONDecodeError:
        st.error(f"Failed to parse AI response.\nRaw output:\n{content}")
        return []
    except Exception as e:
        st.error(f"OpenAI API error: {e}")
        return []

# --- Process file and framework ---
if uploaded_file:
    transcript_text = extract_text(uploaded_file)
    if not transcript_text:
        st.error("Transcript file is empty or could not be read.")
    else:
        st.subheader("Transcript Preview")
        st.text_area("Transcript", transcript_text, height=200)

        # Load framework if provided
        framework_text = None
        if framework_file:
            raw = extract_text(framework_file)
            if framework_file.name.endswith(".json"):
                try:
                    framework_json = json.loads(raw)
                    framework_text = json.dumps(framework_json)
                except json.JSONDecodeError:
                    st.warning("Framework file is not valid JSON. Using raw text.")
                    framework_text = raw
            else:
                framework_text = raw

        if st.button("Generate AI Codes and Highlights"):
            with st.spinner("Analyzing with AI..."):
                results = ai_generate_codes(transcript_text, analysis_mode, research_questions, framework_text)
            if results:
                # Filter valid records
                filtered = [r for r in results if all(k in r for k in ("Text","Code","Theme"))]
                if not filtered:
                    st.warning("AI did not return expected fields.")
                df = pd.DataFrame(filtered)
                st.subheader("Coded Segments")
                st.dataframe(df)

                # Highlight transcript
                st.subheader("Highlighted Transcript")
                highlights = transcript_text
                used = set()
                for row in filtered:
                    txt = row['Text']
                    tag = f"🔹[{row['Code']}]"
                    if txt and txt not in used:
                        highlights = highlights.replace(txt, f"{tag} {txt}", 1)
                        used.add(txt)
                st.text_area("Highlights", highlights, height=300)

                # CSV download
                csv = df.to_csv(index=False)
                st.download_button(
                    "Download Code Table as CSV", csv, "coded_segments.csv", "text/csv"
                )

                # Word report
                doc = Document()
                doc.add_heading('Qualitative Analysis Report', 0)
                doc.add_paragraph('Research Questions: ' + research_questions)
                doc.add_heading('Codes and Themes', level=1)
                table = doc.add_table(rows=1, cols=3)
                hdr = table.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text = 'Text','Code','Theme'
                for _, r in df.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(r['Text'])
                    row_cells[1].text = str(r['Code'])
                    row_cells[2].text = str(r['Theme'])
                doc_stream = BytesIO()
                doc.save(doc_stream)
                doc_stream.seek(0)
                st.download_button(
                    "Download Word Report", doc_stream,
                    "report.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # PDF report
                pdf_stream = BytesIO()
                pdf = SimpleDocTemplate(pdf_stream, pagesize=letter)
                styles = getSampleStyleSheet()
                elems = [Paragraph('Qualitative Analysis Report', styles['Title']),
                         Paragraph('Research Questions: ' + research_questions, styles['Normal']), Spacer(1,12)]
                table_data = [['Text','Code','Theme']] + df.values.tolist()
                elems.append(Table(table_data))
                pdf.build(elems)
                pdf_stream.seek(0)
                st.download_button(
                    "Download PDF Report", pdf_stream, "report.pdf", "application/pdf"
                )
            else:
                st.warning("No results returned by the AI. Please check your input and try again.")
        else:
            st.warning("No results returned by the AI. Please check your input and try again.")